from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("期末复习-数据类型基础知识点与讲解.docx")


def set_east_asian_font(run, font_name="Microsoft YaHei"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.25


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for line in code.splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(31, 77, 120)
    return p


def add_tip(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.4)
    cell = table.cell(0, 0)
    cell.width = Inches(6.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "：")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    set_east_asian_font(r)
    r = p.add_run(text)
    set_east_asian_font(r)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        set_east_asian_font(run)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_east_asian_font(run)
    doc.add_paragraph()
    return table


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_east_asian_font(r)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_east_asian_font(r)
    return p


doc = Document()
style_doc(doc)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(3)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Python 数据类型基础知识点与讲解")
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = RGBColor(11, 37, 69)
set_east_asian_font(r)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(12)
r = sub.add_run("依据：第2周基本数据类型课件与 notebook、期末复习汇总框架、Paper2 示范试题")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(85, 85, 85)
set_east_asian_font(r)

add_tip(
    doc,
    "复习定位",
    "本讲义先整理数据类型的基础部分。主线是数字、布尔与运算、字符串；同时补上期末复习框架中会继续出现的组合数据类型入口，方便后续接列表、元组、集合、字典。",
)

doc.add_heading("1. 数据类型总览：先知道 Python 在表示什么", level=1)
doc.add_paragraph(
    "数据类型决定了一个值能表示什么、能参与什么运算、运算结果会变成什么类型。考试里常见的不是背概念，而是给一小段代码，让你判断输出、类型或是否报错。"
)
add_table(
    doc,
    ["类别", "常见类型", "核心用途", "期末常见考法"],
    [
        ("数字", "int, float, complex", "表示整数、小数、复数并参与数值计算", "进制、精度、类型转换、运算符优先级"),
        ("布尔", "bool", "表示条件成立与否：True / False", "关系运算、逻辑运算、if 条件判断"),
        ("字符串", "str", "表示文本，是不可变的字符序列", "索引、切片、拼接、format 格式化"),
        ("组合类型", "list, tuple, set, dict", "把多个数据组织起来", "可变/不可变、索引、键值对、遍历"),
    ],
    [0.9, 1.35, 2.05, 2.1],
)

doc.add_heading("2. 数字类型：int、float、complex", level=1)
doc.add_heading("2.1 整数 int", level=2)
doc.add_paragraph("整数 int 与数学中的整数类似，在 Python 中通常不用担心取值范围限制。因此课件中用 pow(2, 1000) 这类大数来展示 Python 处理大整数的能力。")
bullet(doc, "十进制直接写：10、-2010。")
bullet(doc, "二进制以 0b 或 0B 开头，例如 0b11 等于十进制 3。")
bullet(doc, "八进制以 0o 或 0O 开头，例如 0o11 等于十进制 9。")
bullet(doc, "十六进制以 0x 或 0X 开头，例如 0x11 等于十进制 17。")
add_code(doc, "print(0b11)   # 3\nprint(0o11)   # 9\nprint(0x11)   # 17\nprint(pow(2, 10))  # 1024")
add_tip(doc, "易错点", "数字开头的 0b、0o、0x 不是字符串装饰，而是在告诉 Python 按二进制、八进制、十六进制解释这个整数。")

doc.add_heading("2.2 浮点数 float", level=2)
doc.add_paragraph("浮点数用于表示带小数的数，也可以用科学计数法表示。它适合近似计算，但不是无限精确的数学实数。")
bullet(doc, "普通小数：0.12、-77.0、-2.17。")
bullet(doc, "科学计数法：aEb 表示 a * 10**b，例如 4.3e2 等于 430.0。")
bullet(doc, "浮点数有精度限制，过大可能 overflow，过小的差异可能被舍入。")
add_code(doc, "i = 1.0 + pow(2, -5000)\nprint(i)  # 很可能仍然显示 1.0，因为差异太小")
add_tip(doc, "Paper2 对应考法", "示范试题里用自然常数 e 的近似公式考察浮点精度：N 先变大时估计更准，但 N 极大时 1 + 1/N 在浮点表示中会被舍入成 1.0，最后结果反而变成 1.0。")

doc.add_heading("2.3 复数 complex", level=2)
doc.add_paragraph("Python 用 j 表示虚数单位。复数可以通过 .real 和 .imag 分别访问实部和虚部。")
add_code(doc, "z = 1.0 + 1.0j\nprint(z.real, z.imag)  # 1.0 1.0\nprint(z ** 2)          # 2j")

doc.add_heading("3. 运算、类型变宽与转换", level=1)
doc.add_heading("3.1 运算结果的类型", level=2)
doc.add_paragraph("不同数字类型混合运算时，结果通常会向“更宽”的类型转换：int -> float -> complex。")
add_table(
    doc,
    ["表达式", "结果示意", "原因"],
    [
        ("123 + 4", "127，int", "整数加整数，结果仍是整数"),
        ("123 + 4.0", "127.0，float", "int 与 float 混合，结果变成 float"),
        ("5 / 2", "2.5，float", "/ 总是返回 float"),
        ("5 // 2", "2，int", "// 是整除，返回商的整数部分"),
        ("1 + 2j + 3", "(4+2j)，complex", "复数参与后结果为 complex"),
    ],
    [1.35, 1.45, 3.5],
)
doc.add_heading("3.2 显式类型转换", level=2)
bullet(doc, "float(5) 得到 5.0，相当于补上小数部分。")
bullet(doc, "int(5.6) 得到 5，注意不是四舍五入，而是直接截去小数部分。")
bullet(doc, "str(123) 得到 '123'，从此以后 + 表示字符串拼接，而不是数字加法。")
add_code(doc, "a = 123\nb = 123e-5\nprint(a + b)          # 123.00123，数字加法\nprint(str(a) + str(b)) # '1230.00123'，字符串拼接")

doc.add_heading("3.3 运算符优先级", level=2)
doc.add_paragraph("优先级决定一行表达式先算什么。复习时最重要的是：括号最高，幂运算 ** 需要特别小心。")
numbered(doc, "括号 ()：最高优先级，复杂表达式建议主动加括号。")
numbered(doc, "幂运算 **：例如 2 ** 2 ** 3 等价于 2 ** (2 ** 3)，结果是 256。")
numbered(doc, "乘除取余：*, /, %, //。")
numbered(doc, "加减：+, -。")
numbered(doc, "比较运算：==, !=, >, <, >=, <=。")
numbered(doc, "逻辑运算：not 高于 and，and 高于 or。")
add_tip(doc, "考试建议", "遇到 2 ** 2 ** 3、x % 2 == 0 and ... 这类题，不要凭感觉，从优先级或加括号后的结构一步步算。")

doc.add_heading("4. 布尔值、关系运算与逻辑运算", level=1)
doc.add_paragraph("关系运算的结果是布尔值 bool，即 True 或 False。逻辑运算用于把多个条件组合起来。")
add_table(
    doc,
    ["运算", "含义", "记忆方式"],
    [
        ("and", "与", "全真才真"),
        ("or", "或", "全假才假"),
        ("not", "非", "真变假、假变真"),
    ],
    [1.0, 1.0, 4.2],
)
add_code(doc, "y = 2000\nis_leap = (y % 4 == 0 and y % 100 != 0) or (y % 400 == 0)\nprint(is_leap)  # True")
doc.add_paragraph("闰年判断是课件中的典型例子：它把取余、关系运算和逻辑运算组合到一起，很适合训练“按条件拆分”的能力。")

doc.add_heading("5. 字符串 str：不可变的字符序列", level=1)
doc.add_heading("5.1 创建与转义", level=2)
doc.add_paragraph("字符串是用单引号、双引号或三引号括起来的字符序列。三引号可以跨多行。反斜杠 \\ 用作转义符。")
bullet(doc, "输出引号：\"带\\\"的字符串\"。")
bullet(doc, "输出反斜杠：\"带\\\\的字符串\"。")
bullet(doc, "换行与制表：\\n 表示换行，\\t 表示制表符。")

doc.add_heading("5.2 索引与切片", level=2)
doc.add_paragraph("字符串是序列，索引从 0 开始；负索引从右往左数，最右边是 -1。切片格式是 s[start:end:step]，包含 start，不包含 end。")
add_table(
    doc,
    ["写法", "含义", "例子"],
    [
        ("s[i]", "取第 i 个字符", "s[0] 取第一个字符"),
        ("s[-1]", "取最后一个字符", "适合从右侧定位"),
        ("s[a:b]", "从 a 到 b-1", "s[0:3] 取前 3 个字符"),
        ("s[a:]", "从 a 到结尾", "s[2:]"),
        ("s[:b]", "从开头到 b-1", "s[:4]"),
        ("s[a:b:k]", "按步长 k 切片", "s[1::2] 从索引 1 开始隔一个取一个"),
    ],
    [1.2, 2.2, 2.8],
)
add_code(doc, "s = '白日依山尽，黄河入海流。'\nprint(s[1::2])  # 从索引1开始，每隔1个字符取一次")
add_tip(doc, "Paper2 对应考法", "示范题直接考 s[1::2] 的输出。做这类题先给每个字符标索引，再按 start、end、step 取值；不要把 end 当成会被取到的位置。")

doc.add_heading("5.3 字符串运算与常用方法", level=2)
bullet(doc, "s1 + s2：拼接两个字符串。")
bullet(doc, "s * n：把字符串重复 n 次。")
bullet(doc, "len(s)：返回字符串长度；中文字符按字符数计算。")
bullet(doc, "replace(old, new)：返回替换后的新字符串。")
bullet(doc, "upper() / lower()：返回大小写转换后的新字符串。")
doc.add_paragraph("关键点：字符串不可变。方法不会修改原字符串，而是返回一个新字符串。如果没有用变量接住结果，原字符串不变。")
add_code(doc, "s = 'Hello John'\nt = s.replace('John', 'Python')\nprint(s)  # Hello John\nprint(t)  # Hello Python")

doc.add_heading("5.4 format 格式化", level=2)
doc.add_paragraph("format 用于把数据填入模板字符串，常见于日志、进度条、对齐输出等题型。")
add_code(doc, "pattern = '{0}: 计算机{1}的CPU占有率为{2}%'\nprint(pattern.format('2026-06-15 12:00', 'Python', 10))\n\nprint('{0:>3.0f}%'.format(8.6))  # 宽度3，右对齐，保留0位小数")
bullet(doc, "花括号中的 0、1、2 表示参数位置。")
bullet(doc, ":>3.0f 表示右对齐、宽度 3、浮点数保留 0 位小数。")
bullet(doc, "\\r 配合 print(..., end='') 可以做单行刷新的进度条。")

doc.add_heading("6. 与期末复习框架衔接：组合数据类型入口", level=1)
doc.add_paragraph("期末复习汇总从组合数据类型展开，所以基础数据类型复习完以后，要自然过渡到 list、tuple、set、dict。这里先放最核心的判断标准。")
add_table(
    doc,
    ["类型", "是否有序", "是否可变", "典型考点"],
    [
        ("list 列表", "有序", "可变", "索引、切片、append、嵌套列表修改"),
        ("tuple 元组", "有序", "不可变", "元组本身不可改，但里面若含列表，列表对象可改"),
        ("set 集合", "无序", "可变", "去重、成员判断、不能索引"),
        ("dict 字典", "按键访问", "可变", "键唯一、get、keys/values/items、遍历"),
    ],
    [1.2, 1.15, 1.15, 2.7],
)
add_tip(doc, "Paper2 对应考法", "示范题中出现了 t[-1][-1] += 7、a[-1].append(1) 这类题。核心是先判断最外层对象能不能改，再判断被取出的内部对象是什么类型。列表可变，元组不可变；但元组中保存的列表仍然可以被列表方法修改。")

doc.add_heading("7. 最小复习清单", level=1)
for item in [
    "能写出 int、float、complex、str、bool 的例子，并能用 type(x) 判断类型。",
    "能解释 0b、0o、0x 分别表示什么进制。",
    "能区分 / 与 //，能解释 int(5.6) 为什么是 5。",
    "能解释浮点数为什么会有精度误差，而不是把 Python 当成算错。",
    "能按优先级计算 2 ** 2 ** 3、带 and/or/not 的表达式。",
    "能熟练做字符串索引、负索引、切片，尤其是 s[start:end:step]。",
    "能区分数字加法与字符串拼接。",
    "能看懂 format 中的位置参数和简单格式控制。",
    "能初步判断 list、tuple、set、dict 的可变性与访问方式。",
]:
    bullet(doc, item)

doc.add_heading("8. 练习建议", level=1)
doc.add_paragraph("建议你复习时不要只看文字，而是打开 Jupyter Notebook 逐段运行。每遇到一个例子，先手算输出，再运行验证。这样最接近期末代码读输出题的要求。")
numbered(doc, "先做数字：进制、pow、/、//、%、类型转换。")
numbered(doc, "再做字符串：索引、切片、replace、format。")
numbered(doc, "最后做组合类型：list/tuple 的嵌套修改、dict 的键值访问。")

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Python 与人工智能期末复习 | 数据类型基础")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(85, 85, 85)
set_east_asian_font(run)

doc.save(OUT)
print(OUT.resolve())
