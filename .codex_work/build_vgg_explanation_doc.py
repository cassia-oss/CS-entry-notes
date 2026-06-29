import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NB_PATH = Path("python与人工智能-代码/VGG/VGG代码.ipynb")
OUT_PATH = Path("VGG代码逐行解释文档.docx")


MODULES = {
    0: ("模块一：导入依赖库", "这一部分准备 PyTorch、神经网络模块、预训练权重下载函数和类型注解工具。"),
    1: ("模块二：公开接口与预训练模型地址", "这一部分定义当前文件对外提供哪些 VGG 构造函数，以及不同 VGG 版本的 ImageNet 预训练权重下载地址。"),
    2: ("模块三：VGG 主体类", "这一部分定义 VGG 网络的整体结构：特征提取层 features、自适应池化 avgpool、分类器 classifier、前向传播 forward 和参数初始化。"),
    3: ("模块四：根据配置表生成卷积层", "这一部分把配置列表 cfg 转换成真正的 nn.Sequential 网络层。数字代表卷积层输出通道数，'M' 代表最大池化层。"),
    4: ("模块五：VGG 网络配置表", "这一部分用 A、B、D、E 四套配置表示 VGG11、VGG13、VGG16、VGG19 的卷积结构。"),
    5: ("模块六：统一构造 VGG 模型的内部函数", "这一部分负责根据模型名称、配置编号、是否使用 BatchNorm、是否加载预训练权重来创建具体模型。"),
    6: ("模块七：不同 VGG 版本的工厂函数", "这一部分提供 vgg11、vgg13、vgg16、vgg19 及其 BatchNorm 版本。用户通常不直接调用 _vgg，而是调用这些函数。"),
}


def set_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 15, "2E74B5", 12, 7),
        ("Heading 2", 12.5, "2E74B5", 9, 5),
        ("Heading 3", 11.5, "1F4D78", 7, 4),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15


def para(doc, text, style=None, size=None, color=None, bold=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold)
    return p


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(9.7)
    cell = table.cell(0, 0)
    cell.width = Inches(9.7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    r = p.add_run(label + "：")
    set_font(r, bold=True, color="1F4D78")
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph()


def code_cell_text(cell):
    return "".join(cell.get("source", [])).rstrip("\n")


def clean_code(line):
    return line.rstrip("\n")


def explain(line, cell_idx, prev_nonblank=""):
    stripped = line.strip()
    if not stripped:
        return "空行，用来分隔代码块，提高可读性，不影响程序运行。"
    if stripped.startswith("#"):
        return "注释行，说明作者对这一段代码的理解或意图，Python 运行时会忽略。"
    if stripped.startswith("import torch"):
        return "导入 PyTorch 主库。后面会用 torch.Tensor、torch.flatten 等张量相关功能。"
    if stripped == "import torch.nn as nn":
        return "导入 PyTorch 的神经网络模块，并简写为 nn。卷积层、池化层、全连接层、激活函数都来自这里。"
    if "load_state_dict_from_url" in stripped and stripped.startswith("from torch.hub"):
        return "从 torch.hub 导入下载预训练权重的函数。pretrained=True 时会用它从网址下载模型参数。"
    if stripped.startswith("from typing"):
        return "导入类型注解工具。Union、List、Dict、Any、cast 主要帮助说明变量或参数应该是什么类型。"
    if stripped.startswith("__all__"):
        return "__all__ 定义这个文件对外暴露的名字。别人 from xxx import * 时，只会导入列表里的对象。"
    if re.match(r"^'.+',?$", stripped):
        return "这是 __all__ 列表中的一个字符串元素，表示一个可被外部导入的类或函数名。"
    if stripped == "]":
        return "结束 __all__ 列表。"
    if stripped.startswith("model_urls"):
        return "创建字典 model_urls，用模型名称映射到对应预训练权重文件的下载地址。"
    if re.match(r"^'vgg.*\.pth',?$", stripped):
        return "字典中的一项：键是模型版本名，值是该版本 ImageNet 预训练参数的下载链接。"
    if stripped == "}":
        return "结束字典定义。"
    if stripped.startswith("class VGG"):
        return "定义 VGG 类，并继承 nn.Module。继承后它才能被 PyTorch 当作神经网络模型使用。"
    if stripped.startswith("def __init__"):
        return "定义构造函数。创建 VGG 对象时，会先运行这里来搭建网络结构。"
    if stripped == "self,":
        return "self 表示当前正在创建的 VGG 实例。类里的方法都通过 self 访问对象自身的属性。"
    if stripped.startswith("features: nn.Module"):
        return "features 参数表示前面的卷积和池化特征提取部分，类型注解为 nn.Module。"
    if stripped.startswith("num_classes"):
        return "num_classes 表示最终分类类别数，默认 1000，对应 ImageNet 的 1000 类。"
    if stripped.startswith("init_weights"):
        return "init_weights 控制是否执行自定义参数初始化。默认 True，表示从零训练时自动初始化。"
    if stripped.startswith(") -> None"):
        return "类型注解：这个函数不返回有意义的值。构造函数通常只负责初始化对象。"
    if stripped.startswith("super("):
        return "调用父类 nn.Module 的初始化逻辑。自定义 PyTorch 模型时这一步很重要。"
    if stripped.startswith("self.features"):
        return "把传入的卷积/池化模块保存为模型的 features 部分。forward 中会先调用它提取图像特征。"
    if stripped.startswith("self.avgpool"):
        return "定义自适应平均池化层，把特征图统一变成 7 x 7，方便后面接固定长度的全连接层。"
    if stripped.startswith("self.classifier"):
        return "定义分类器部分，用 nn.Sequential 把多个全连接、ReLU、Dropout 层按顺序串起来。"
    if "nn.Linear(512 * 7 * 7, 4096)" in stripped:
        return "第一个全连接层。输入维度是 512 个通道乘 7 x 7，输出 4096 维特征。"
    if stripped == "nn.ReLU(True),":
        return "ReLU 激活函数。True 表示 inplace=True，尽量在原张量上操作以节省显存。"
    if stripped == "nn.Dropout(),":
        return "Dropout 层。训练时随机丢弃一部分神经元，降低过拟合风险。"
    if "nn.Linear(4096, 4096)" in stripped:
        return "第二个全连接层，把 4096 维特征继续映射到 4096 维。"
    if "nn.Linear(4096, num_classes)" in stripped:
        return "最后一个全连接层，把 4096 维映射到类别数 num_classes，输出每一类的得分。"
    if stripped in [")", ")#这是最后统一要通过的几层"]:
        return "结束 nn.Sequential 或函数调用的括号。"
    if stripped.startswith("if init_weights"):
        return "如果需要初始化权重，就执行下面的 _initialize_weights 方法。"
    if stripped.startswith("self._initialize_weights"):
        return "调用本类中定义的参数初始化函数，给卷积层、BN 层和全连接层设置初始值。"
    if stripped.startswith("def forward"):
        return "定义前向传播。输入 x 进入模型后，每一步如何计算输出都写在这里。"
    if stripped == "x = self.features(x)":
        return "让输入图像先经过卷积层、ReLU、池化层组成的特征提取网络。"
    if stripped == "x = self.avgpool(x)":
        return "把特征图池化到固定大小 7 x 7。这样无论前面空间尺寸如何，后面全连接层输入维度都稳定。"
    if stripped.startswith("x = torch.flatten"):
        return "从第 1 个维度开始展平。保留 batch 维度，把每张图的 512 x 7 x 7 特征拉成一维向量。"
    if stripped == "x = self.classifier(x)":
        return "把展平后的特征送入全连接分类器，得到每个类别的预测分数。"
    if stripped == "return x":
        return "返回模型输出。对分类任务来说，这通常是还没有经过 softmax 的 logits。"
    if stripped.startswith("def _initialize_weights"):
        return "定义权重初始化方法。不同层类型适合不同的初始化方式。"
    if stripped.startswith("for m in self.modules"):
        return "遍历模型中的所有子模块，比如 Conv2d、BatchNorm2d、Linear。"
    if "isinstance(m, nn.Conv2d)" in stripped:
        return "判断当前模块是否是二维卷积层。卷积层使用 Kaiming 初始化更适合 ReLU 网络。"
    if "kaiming_normal_" in stripped:
        return "用 Kaiming normal 方法初始化卷积核权重，fan_out 模式常用于卷积网络。"
    if stripped.startswith("if m.bias is not None"):
        return "如果这一层有 bias 参数，才对 bias 初始化；有些层可能没有 bias。"
    if "constant_(m.bias, 0)" in stripped:
        return "把 bias 初始化为 0。"
    if "isinstance(m, nn.BatchNorm2d)" in stripped:
        return "判断当前模块是否是 BatchNorm2d 层。"
    if "constant_(m.weight, 1)" in stripped:
        return "把 BatchNorm 的缩放参数 gamma 初始化为 1，使其一开始不改变特征尺度。"
    if "isinstance(m, nn.Linear)" in stripped:
        return "判断当前模块是否是全连接层。"
    if "normal_(m.weight, 0, 0.01)" in stripped:
        return "用均值 0、标准差 0.01 的正态分布初始化全连接层权重。"
    if stripped.startswith("def make_layers"):
        return "定义辅助函数：把配置列表 cfg 转换成真正的卷积、BN、ReLU、池化层序列。"
    if stripped.startswith("layers: List"):
        return "创建空列表 layers，用来逐步收集网络层。类型注解说明列表元素是 nn.Module。"
    if stripped.startswith("in_channels = 3"):
        return "输入图像默认有 3 个通道，即 RGB 三通道。第一层卷积的输入通道数就是 3。"
    if stripped.startswith("for v in cfg"):
        return "逐个读取配置表中的元素。数字表示卷积输出通道数，'M' 表示池化层。"
    if stripped.startswith("if v == 'M'"):
        return "如果当前配置是 'M'，就添加最大池化层，而不是卷积层。"
    if "nn.MaxPool2d" in stripped:
        return "添加 2 x 2 最大池化，stride=2 表示宽高减半，用于降低特征图尺寸。"
    if stripped.startswith("else"):
        return "如果当前配置不是 'M'，说明它是一个数字，要创建卷积层。"
    if stripped.startswith("v = cast"):
        return "把 v 告诉类型检查器：这里可以把它当作 int 使用。运行时主要是类型说明作用。"
    if stripped.startswith("conv2d = nn.Conv2d"):
        return "创建二维卷积层。输入通道是 in_channels，输出通道是 v，卷积核 3 x 3，padding=1 保持尺寸基本不变。"
    if stripped.startswith("if batch_norm"):
        return "如果用户选择 BatchNorm 版本，就在卷积层后加入 BatchNorm2d。"
    if "nn.BatchNorm2d(v)" in stripped:
        return "添加卷积层、BatchNorm 和 ReLU。BatchNorm 的通道数要等于当前卷积输出通道 v。"
    if "layers += [conv2d, nn.ReLU" in stripped:
        return "不使用 BatchNorm 时，只添加卷积层和 ReLU 激活函数。"
    if stripped.startswith("in_channels = v"):
        return "更新下一层的输入通道数。当前卷积输出 v 个通道，下一层就要接收 v 个通道。"
    if stripped.startswith("return nn.Sequential"):
        return "用 *layers 把列表拆开，传给 nn.Sequential，形成按顺序执行的网络模块。"
    if stripped.startswith("cfgs: Dict"):
        return "定义配置字典 cfgs。键 A/B/D/E 表示不同 VGG 深度，值是卷积/池化结构列表。"
    if stripped.startswith("'A':"):
        return "配置 A，对应 VGG11。数字是卷积层输出通道，'M' 是最大池化。"
    if stripped.startswith("'B':"):
        return "配置 B，对应 VGG13，比 A 多了一些卷积层。"
    if stripped.startswith("'D':"):
        return "配置 D，对应 VGG16，是经典 VGG16 的卷积结构。"
    if stripped.startswith("'E':"):
        return "配置 E，对应 VGG19，比 VGG16 更深。"
    if stripped.startswith("def _vgg"):
        return "定义内部构造函数 _vgg。它统一处理不同 VGG 版本的创建和预训练权重加载。"
    if stripped.startswith("if pretrained"):
        return "如果要求加载预训练模型，就进入对应逻辑。"
    if "kwargs['init_weights'] = False" in stripped:
        return "加载预训练权重时不需要随机初始化，因为后面会用下载好的参数覆盖。"
    if stripped.startswith("model = VGG"):
        return "先用 cfgs[cfg] 生成 features，再创建 VGG 对象。**kwargs 可传入 num_classes 等额外参数。"
    if stripped.startswith("state_dict = load_state_dict_from_url"):
        return "从 model_urls[arch] 指定的网址下载预训练权重，得到 state_dict 参数字典。"
    if stripped.startswith("progress=progress"):
        return "把 progress 参数传给下载函数，用来控制是否显示下载进度条。"
    if stripped.startswith("model.load_state_dict"):
        return "把下载到的参数加载进模型，使模型拥有 ImageNet 预训练能力。"
    if stripped == "return model":
        return "返回创建好的 VGG 模型对象。"
    if stripped.startswith("def vgg"):
        name = stripped.split("(")[0].replace("def ", "")
        return f"定义工厂函数 {name}。用户调用它即可得到对应版本的 VGG 模型。"
    if stripped.startswith("r\"\"\"") or stripped == '"""':
        return "文档字符串 docstring 的开始或结束，用来说明函数用途和参数含义。"
    if "VGG 11-layer" in stripped:
        return "文档字符串内容：说明这是 11 层 VGG，对应配置 A。"
    if "VGG 13-layer" in stripped:
        return "文档字符串内容：说明这是 13 层 VGG，对应配置 B。"
    if "VGG 16-layer" in stripped:
        return "文档字符串内容：说明这是 16 层 VGG，对应配置 D。"
    if "VGG 19-layer" in stripped:
        return "文档字符串内容：说明这是 19 层 VGG，对应配置 E。"
    if "Very Deep Convolutional Networks" in stripped:
        return "文档字符串内容：引用 VGG 原论文。VGG 的核心思想是使用很深的 3 x 3 卷积网络。"
    if stripped.startswith("Args"):
        return "文档字符串内容：下面开始解释函数参数。"
    if stripped.startswith("pretrained"):
        return "文档字符串内容：pretrained=True 时返回在 ImageNet 上预训练过的模型。"
    if stripped.startswith("progress"):
        return "文档字符串内容：progress=True 时下载权重会显示进度。"
    if stripped.startswith("return _vgg"):
        args = stripped[stripped.find("(") + 1 : stripped.rfind(")")]
        return f"调用统一构造函数 _vgg，传入版本信息 {args}，返回具体的 VGG 模型。"
    return "这一行属于当前模块的语法结构或参数说明，作用需要结合上下文理解：它参与定义网络、配置或函数调用。"


def add_code_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.6, 4.15, 4.95]
    headers = ["行号", "原代码", "逐行解释"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True, size=9.5)
    for line_no, code, desc in rows:
        cells = table.add_row().cells
        values = [str(line_no), code if code.strip() else "(空行)", desc]
        for i, val in enumerate(values):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                set_font(r, size=8.5)
            elif i == 1:
                r = p.add_run(val)
                set_font(r, name="Consolas", size=7.8, color="0B2545")
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            else:
                r = p.add_run(val)
                set_font(r, size=8.8)
    doc.add_paragraph()


def main():
    nb = json.loads(NB_PATH.read_text(encoding="utf-8"))
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("VGG代码.ipynb 分模块逐行解释")
    set_font(r, size=20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("来源文件：python与人工智能-代码/VGG/VGG代码.ipynb")
    set_font(r, size=9.5, color="555555")

    add_note(
        doc,
        "阅读主线",
        "这份代码的目标是定义 VGG 系列图像分类模型。先定义卷积特征提取部分，再定义全连接分类器；不同 VGG 版本的差别主要来自 cfgs 配置表。",
    )

    doc.add_heading("整体流程图", level=1)
    for text in [
        "导入 PyTorch 与类型工具。",
        "定义可导出的函数名与预训练权重地址。",
        "定义 VGG 类：features -> avgpool -> flatten -> classifier。",
        "用 make_layers 根据 cfgs 生成卷积/池化网络。",
        "用 _vgg 统一创建模型，并在需要时加载预训练参数。",
        "用 vgg11/vgg13/vgg16/vgg19 等工厂函数提供简单调用入口。",
    ]:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        set_font(r)

    global_line = 1
    for cell_idx, cell in enumerate(nb["cells"]):
        if cell.get("cell_type") != "code":
            continue
        source = code_cell_text(cell)
        if not source.strip():
            continue
        heading, summary = MODULES.get(cell_idx, (f"模块：代码单元 {cell_idx}", "该单元包含补充代码。"))
        doc.add_heading(heading, level=1)
        para(doc, summary)

        rows = []
        prev = ""
        for local_line, line in enumerate(source.splitlines(), start=1):
            code = clean_code(line)
            desc = explain(code, cell_idx, prev)
            rows.append((global_line, code, desc))
            if code.strip():
                prev = code.strip()
            global_line += 1
        add_code_table(doc, rows)

    doc.add_heading("关键概念补充", level=1)
    add_note(
        doc,
        "VGG 的核心",
        "VGG 网络大量使用 3 x 3 卷积和 2 x 2 最大池化。卷积层逐步增加通道数，池化层逐步缩小图像宽高，最后用全连接层完成分类。",
    )
    add_note(
        doc,
        "Batch Normalization",
        "带 _bn 的版本会在卷积层后加入 BatchNorm2d。它可以让训练更稳定，也常常能加快收敛。",
    )
    add_note(
        doc,
        "pretrained",
        "pretrained=True 表示加载在 ImageNet 上训练好的权重。学习代码结构时可以先设为 False，避免下载模型文件；真正迁移学习时再考虑设为 True。",
    )
    add_note(
        doc,
        "一个小问题",
        "notebook 最后重复定义了一次 vgg11。Python 会用后出现的定义覆盖前一个同名函数；这里功能基本相同，所以影响不大，但写项目时应避免重复定义。",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Python 与人工智能 | VGG 代码逐行解释")
    set_font(r, size=8.5, color="555555")

    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    main()
